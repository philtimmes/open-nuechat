"""
RAG Service with FAISS GPU (ROCm) and sentence-transformers

Uses:
- sentence-transformers for embeddings (ROCm/GPU accelerated via PyTorch)
- FAISS GPU for vector indexing and search (ROCm)
- IVF-PQ index for efficient large dataset handling

Index Strategy:
- IVF (Inverted File) for fast approximate search
- PQ (Product Quantization) for memory efficiency
- Supports millions of vectors efficiently
"""
import os
import pickle
import asyncio
import re
from typing import List, Dict, Any, Optional, Tuple, Set
from pathlib import Path
import numpy as np
from concurrent.futures import ThreadPoolExecutor
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, delete
import logging

import faiss

from app.core.config import settings
from app.models.models import Document, DocumentChunk, User, KnowledgeStore, KnowledgeStoreShare

logger = logging.getLogger(__name__)

# Thread pool for CPU-bound operations
_executor = ThreadPoolExecutor(max_workers=4)

# Cache for debug_rag setting (refreshed on each search)
_debug_rag_cache: Dict[str, bool] = {"enabled": False}


def _parse_keywords_string(keywords_str: str) -> Tuple[List[str], List[str]]:
    """
    Parse a keyword string into phrases and keywords.
    
    NC-0.8.0.1.1: Supports:
    - "exact phrases" in quotes
    - comma-separated keywords
    
    Example: '"alien spacecraft", UFO, glimmer, "unidentified object"'
    Returns: (phrases=["alien spacecraft", "unidentified object"], keywords=["ufo", "glimmer"])
    """
    if not keywords_str or not keywords_str.strip():
        return [], []
    
    phrases = []
    keywords = []
    
    # Extract quoted phrases first
    phrase_pattern = r'"([^"]+)"'
    found_phrases = re.findall(phrase_pattern, keywords_str)
    phrases = [p.strip().lower() for p in found_phrases if p.strip()]
    
    # Remove quoted phrases from string
    remaining = re.sub(phrase_pattern, '', keywords_str)
    
    # Split remaining by comma and clean up
    for kw in remaining.split(','):
        kw = kw.strip().lower()
        if kw:
            keywords.append(kw)
    
    return phrases, keywords


def _check_keywords_match(
    query: str,
    history: str,
    phrases: List[str],
    keywords: List[str],
    match_mode: str = 'any'
) -> bool:
    """
    Check if query/history matches required keywords.
    
    NC-0.8.0.1.1 match modes:
    - 'any': Any phrase OR any keyword matches
    - 'all': ALL phrases AND ALL keywords must match
    - 'mixed': ALL phrases must match, but only ANY keyword needs to match
    
    Args:
        query: User's current query (lowercased)
        history: Recent conversation history (lowercased)
        phrases: Exact phrases that must appear
        keywords: Individual keywords
        match_mode: 'any', 'all', or 'mixed'
    
    Returns:
        True if keywords match according to mode
    """
    if not phrases and not keywords:
        return True  # No requirements = always match
    
    combined_text = f"{query} {history}".lower()
    
    # Check phrase matches
    phrase_matches = [p in combined_text for p in phrases]
    
    # Check keyword matches
    keyword_matches = [kw in combined_text for kw in keywords]
    
    if match_mode == 'all':
        # All phrases AND all keywords must match
        return all(phrase_matches) and all(keyword_matches) if (phrases or keywords) else True
    
    elif match_mode == 'mixed':
        # All phrases must match, but only need any keyword
        phrases_ok = all(phrase_matches) if phrases else True
        keywords_ok = any(keyword_matches) if keywords else True
        return phrases_ok and keywords_ok
    
    else:  # 'any' (default)
        # Any phrase OR any keyword matches
        return any(phrase_matches) or any(keyword_matches)


async def _is_debug_rag_enabled(db: AsyncSession) -> bool:
    """Check if debug RAG logging is enabled."""
    from app.models.models import SystemSetting
    try:
        result = await db.execute(
            select(SystemSetting).where(SystemSetting.key == "debug_rag")
        )
        setting = result.scalar_one_or_none()
        enabled = setting and setting.value == "true"
        _debug_rag_cache["enabled"] = enabled
        return enabled
    except Exception:
        return _debug_rag_cache.get("enabled", False)


def _log_rag_debug(message: str, **kwargs):
    """Log RAG debug message with optional structured data."""
    if kwargs:
        data_str = ", ".join(f"{k}={v}" for k, v in kwargs.items())
        logger.info(f"[RAG DEBUG] {message} | {data_str}")
    else:
        logger.info(f"[RAG DEBUG] {message}")


class FAISSIndexManager:
    """
    Manages FAISS GPU indexes with IVF-PQ for efficient large-scale search.
    
    Index types by size:
    - < 10K vectors: Flat (exact search)
    - 10K - 1M vectors: IVF-Flat
    - > 1M vectors: IVF-PQ
    """
    
    # Index configuration
    EMBEDDING_DIM = 384  # Default for all-MiniLM-L6-v2
    NLIST_FACTOR = 4  # nlist = NLIST_FACTOR * sqrt(n)
    NPROBE = 32  # Number of clusters to search (tradeoff: speed vs accuracy)
    PQ_M = 48  # Number of subquantizers for PQ (must divide EMBEDDING_DIM)
    PQ_NBITS = 8  # Bits per subquantizer
    
    def __init__(self, index_dir: str = None):
        self.index_dir = Path(index_dir or settings.FAISS_INDEX_DIR)
        self.index_dir.mkdir(parents=True, exist_ok=True)
        
        # GPU resources
        self.gpu_resources = None
        self.gpu_available = False
        self._init_gpu()
        
        # In-memory index cache
        self._indexes: Dict[str, faiss.Index] = {}
        self._id_maps: Dict[str, Dict[int, str]] = {}  # faiss_id -> chunk_id
    
    def _init_gpu(self):
        """Initialize FAISS GPU resources"""
        try:
            self.gpu_resources = faiss.StandardGpuResources()
            self.gpu_available = True
            logger.info("FAISS GPU initialized successfully")
        except Exception as e:
            logger.warning(f"FAISS GPU not available, using CPU: {e}")
            self.gpu_available = False
    
    def _get_index_path(self, index_id: str) -> Tuple[Path, Path]:
        """Get paths for index and ID map files"""
        index_path = self.index_dir / f"{index_id}.faiss"
        map_path = self.index_dir / f"{index_id}.map"
        return index_path, map_path
    
    def _create_index(self, n_vectors: int) -> faiss.Index:
        """
        Create appropriate FAISS index based on dataset size.
        
        Strategy:
        - Flat: Best accuracy, O(n) search, good for < 10K
        - IVF-Flat: Good accuracy, faster search, 10K - 1M
        - IVF-PQ: Memory efficient, fast, 1M+
        """
        dim = self.EMBEDDING_DIM
        
        if n_vectors < 10_000:
            # Small dataset: exact search
            logger.info(f"Creating Flat index for {n_vectors} vectors")
            index = faiss.IndexFlatIP(dim)  # Inner product (cosine with normalized vectors)
        
        elif n_vectors < 1_000_000:
            # Medium dataset: IVF with flat quantizer
            nlist = max(16, int(self.NLIST_FACTOR * np.sqrt(n_vectors)))
            logger.info(f"Creating IVF-Flat index: {n_vectors} vectors, nlist={nlist}")
            
            quantizer = faiss.IndexFlatIP(dim)
            index = faiss.IndexIVFFlat(quantizer, dim, nlist, faiss.METRIC_INNER_PRODUCT)
        
        else:
            # Large dataset: IVF with product quantization
            nlist = max(256, int(self.NLIST_FACTOR * np.sqrt(n_vectors)))
            logger.info(f"Creating IVF-PQ index: {n_vectors} vectors, nlist={nlist}, m={self.PQ_M}")
            
            quantizer = faiss.IndexFlatIP(dim)
            index = faiss.IndexIVFPQ(
                quantizer, dim, nlist, self.PQ_M, self.PQ_NBITS,
                faiss.METRIC_INNER_PRODUCT
            )
        
        return index
    
    def _to_gpu(self, index: faiss.Index) -> faiss.Index:
        """Move index to GPU if available"""
        if self.gpu_available and self.gpu_resources:
            try:
                return faiss.index_cpu_to_gpu(self.gpu_resources, 0, index)
            except Exception as e:
                logger.warning(f"Failed to move index to GPU: {e}")
        return index
    
    def _to_cpu(self, index: faiss.Index) -> faiss.Index:
        """Move index to CPU for saving"""
        if self.gpu_available:
            try:
                return faiss.index_gpu_to_cpu(index)
            except:
                pass
        return index
    
    def build_index(
        self,
        index_id: str,
        embeddings: np.ndarray,
        chunk_ids: List[str],
    ) -> None:
        """
        Build a new FAISS index from embeddings.
        
        Args:
            index_id: Unique identifier for the index (e.g., knowledge_store_id)
            embeddings: numpy array of shape (n, dim) with normalized embeddings
            chunk_ids: List of chunk IDs corresponding to each embedding
        """
        logger.info(f"[FAISS_BUILD] Building index for: {index_id} (type: {type(index_id).__name__})")
        n_vectors = len(embeddings)
        
        if n_vectors == 0:
            logger.warning(f"No vectors to index for {index_id}")
            return
        
        # Ensure embeddings are float32 and normalized
        embeddings = embeddings.astype(np.float32)
        faiss.normalize_L2(embeddings)
        
        # Create index
        index = self._create_index(n_vectors)
        
        # Train if required (IVF indexes need training)
        if hasattr(index, 'train') and not index.is_trained:
            logger.info(f"Training index {index_id} with {n_vectors} vectors")
            # For training, use CPU then move to GPU
            index.train(embeddings)
        
        # Move to GPU for adding vectors
        index = self._to_gpu(index)
        
        # Set search parameters
        if hasattr(index, 'nprobe'):
            index.nprobe = self.NPROBE
        
        # Add vectors
        index.add(embeddings)
        
        # Create ID mapping
        id_map = {i: chunk_id for i, chunk_id in enumerate(chunk_ids)}
        
        # Save to disk (CPU version)
        cpu_index = self._to_cpu(index)
        index_path, map_path = self._get_index_path(index_id)
        
        logger.info(f"[FAISS_BUILD] Saving index to: {index_path}")
        faiss.write_index(cpu_index, str(index_path))
        with open(map_path, 'wb') as f:
            pickle.dump(id_map, f)
        logger.info(f"[FAISS_BUILD] Index saved successfully: {index_path.exists()}")
        
        # Cache GPU version
        self._indexes[index_id] = index
        self._id_maps[index_id] = id_map
        
        logger.info(f"Built index {index_id}: {n_vectors} vectors, saved to {index_path}")
    
    def load_index(self, index_id: str) -> bool:
        """Load index from disk into memory/GPU"""
        logger.info(f"[FAISS_LOAD] Attempting to load index: {index_id}")
        
        if index_id in self._indexes:
            logger.info(f"[FAISS_LOAD] Index {index_id} already cached in memory")
            return True
        
        index_path, map_path = self._get_index_path(index_id)
        logger.info(f"[FAISS_LOAD] Looking for index file: {index_path}")
        
        if not index_path.exists():
            logger.warning(f"[FAISS_LOAD] Index file NOT FOUND: {index_path}")
            # List what files ARE in the directory
            try:
                if self.index_dir.exists():
                    files = list(self.index_dir.glob("*.faiss"))
                    logger.warning(f"[FAISS_LOAD] Available index files: {[f.name for f in files]}")
                else:
                    logger.warning(f"[FAISS_LOAD] Index directory does not exist: {self.index_dir}")
            except Exception as e:
                logger.warning(f"[FAISS_LOAD] Error listing index dir: {e}")
            return False
        
        try:
            # Load from disk
            logger.info(f"[FAISS_LOAD] Loading index from disk: {index_path}")
            index = faiss.read_index(str(index_path))
            
            # Set search parameters
            if hasattr(index, 'nprobe'):
                index.nprobe = self.NPROBE
            
            # Move to GPU
            index = self._to_gpu(index)
            
            # Load ID map
            with open(map_path, 'rb') as f:
                id_map = pickle.load(f)
            
            self._indexes[index_id] = index
            self._id_maps[index_id] = id_map
            
            logger.info(f"Loaded index {index_id}: {index.ntotal} vectors")
            return True
            
        except Exception as e:
            logger.error(f"Failed to load index {index_id}: {e}")
            return False
    
    def search(
        self,
        index_id: str,
        query_embedding: np.ndarray,
        top_k: int = 10,
    ) -> List[Tuple[str, float]]:
        """
        Search index for nearest neighbors.
        
        Args:
            index_id: Index to search
            query_embedding: Query vector (will be normalized)
            top_k: Number of results
            
        Returns:
            List of (chunk_id, score) tuples
        """
        # Ensure index is loaded
        if not self.load_index(index_id):
            logger.warning(f"[FAISS] Index {index_id} not found or failed to load")
            return []
        
        index = self._indexes[index_id]
        id_map = self._id_maps[index_id]
        
        # Prepare query
        query = query_embedding.astype(np.float32).reshape(1, -1)
        
        # Log query info
        logger.debug(f"[FAISS] Search: index={index_id}, index_vectors={index.ntotal}, query_dim={query.shape[1]}, query_norm_before={np.linalg.norm(query):.4f}")
        
        faiss.normalize_L2(query)
        
        logger.debug(f"[FAISS] Query norm after normalize_L2={np.linalg.norm(query):.4f}")
        
        # Search
        scores, indices = index.search(query, min(top_k, index.ntotal))
        
        logger.debug(f"[FAISS] Raw results: indices={indices[0][:5].tolist()}, scores={scores[0][:5].tolist()}")
        
        # Map back to chunk IDs
        results = []
        for score, idx in zip(scores[0], indices[0]):
            if idx >= 0 and idx in id_map:
                results.append((id_map[idx], float(score)))
        
        logger.debug(f"[FAISS] Mapped {len(results)} results, top_scores={[r[1] for r in results[:3]]}")
        
        return results
    
    def delete_index(self, index_id: str) -> None:
        """Delete index from memory and disk"""
        # Remove from cache
        self._indexes.pop(index_id, None)
        self._id_maps.pop(index_id, None)
        
        # Remove from disk
        index_path, map_path = self._get_index_path(index_id)
        index_path.unlink(missing_ok=True)
        map_path.unlink(missing_ok=True)
        
        logger.info(f"Deleted index {index_id}")
    
    def add_vectors(
        self,
        index_id: str,
        embeddings: np.ndarray,
        chunk_ids: List[str],
    ) -> None:
        """Add vectors to existing index (rebuilds for IVF indexes)"""
        # For IVF indexes, we need to rebuild
        # This is because IVF centroids are fixed after training
        # For incremental updates, you'd need a more complex strategy
        
        if not self.load_index(index_id):
            # No existing index, build new one
            self.build_index(index_id, embeddings, chunk_ids)
            return
        
        # Get existing vectors (need to reload from original data)
        # For now, just rebuild - in production, you'd store vectors separately
        logger.info(f"Rebuilding index {index_id} with new vectors")
        self.build_index(index_id, embeddings, chunk_ids)


class RAGService:
    """
    Retrieval Augmented Generation service using FAISS GPU and sentence-transformers.
    
    Features:
    - GPU-accelerated embeddings via PyTorch ROCm
    - FAISS GPU for vector search
    - IVF-PQ indexing for large datasets
    - Async-friendly with thread pool for blocking ops
    """
    
    _model = None
    _model_loaded = False
    _model_load_failed = False  # Track if model loading permanently failed
    _index_manager = None
    _model_lock = None  # Threading lock for model access
    
    @classmethod
    def _get_lock(cls):
        """Get or create the model lock (lazy init to avoid import-time issues)"""
        if cls._model_lock is None:
            import threading
            cls._model_lock = threading.Lock()
        return cls._model_lock
    
    @classmethod
    def reset_model(cls):
        """Reset model state to allow retry of loading"""
        logger.info("[RAG] Resetting model state for retry...")
        with cls._get_lock():
            cls._model = None
            cls._model_loaded = False
            cls._model_load_failed = False
    
    @classmethod
    def get_model_status(cls) -> dict:
        """Get current model status for debugging"""
        return {
            "loaded": cls._model_loaded,
            "failed": cls._model_load_failed,
            "model_exists": cls._model is not None,
        }
    
    @classmethod
    def get_model(cls):
        """Lazy load the embedding model with GPU support (thread-safe)"""
        # Quick check without lock (common case - already loaded)
        if cls._model_load_failed:
            return None
        if cls._model_loaded and cls._model is not None:
            return cls._model
        
        # Acquire lock for model loading
        with cls._get_lock():
            # Double-check after acquiring lock
            if cls._model_load_failed:
                return None
            if cls._model_loaded and cls._model is not None:
                return cls._model
            
            try:
                import torch
                import os
                
                # CRITICAL: Disable accelerate's device_map which causes meta tensor issues
                # Must be set BEFORE importing sentence_transformers/transformers
                os.environ["ACCELERATE_DISABLE_RICH"] = "1"
                os.environ["TRANSFORMERS_NO_ADVISORY_WARNINGS"] = "1"
                # Disable low_cpu_mem_usage globally to prevent meta tensor issues
                os.environ["TRANSFORMERS_LOW_CPU_MEM_USAGE"] = "0"
                os.environ["ACCELERATE_TORCH_DEVICE"] = "cpu"
                
                # NC-0.8.0.7: Always use online mode - offline mode can cause issues with incomplete caches
                os.environ["TRANSFORMERS_OFFLINE"] = "0"
                os.environ["HF_HUB_OFFLINE"] = "0"
                logger.info(f"HuggingFace online mode enabled")
                
                # Temporarily hide GPU to force CPU loading (avoids meta tensor issues)
                original_cuda_visible = os.environ.get("CUDA_VISIBLE_DEVICES", "")
                os.environ["CUDA_VISIBLE_DEVICES"] = ""
                
                try:
                    from sentence_transformers import SentenceTransformer
                    
                    logger.info(f"Loading embedding model '{settings.EMBEDDING_MODEL}' on CPU...")
                    
                    # NC-0.6.50: Fix meta tensor errors by disabling low_cpu_mem_usage
                    model = SentenceTransformer(
                        settings.EMBEDDING_MODEL, 
                        device="cpu",
                        model_kwargs={"low_cpu_mem_usage": False},
                    )
                    
                    # Verify model works
                    logger.info("Testing model with sample embedding...")
                    test_embed = model.encode("test", convert_to_numpy=True)
                    if test_embed is not None and len(test_embed) > 0:
                        logger.info("Model loaded successfully on CPU")
                    else:
                        raise ValueError("Model encode returned empty result")
                        
                finally:
                    # Restore CUDA visibility
                    if original_cuda_visible:
                        os.environ["CUDA_VISIBLE_DEVICES"] = original_cuda_visible
                    else:
                        os.environ.pop("CUDA_VISIBLE_DEVICES", None)
                
                # Now try to move to GPU if available
                target_device = "cuda" if torch.cuda.is_available() else "cpu"
                if target_device == "cuda" and model is not None:
                    try:
                        model = model.to("cuda")
                        logger.info("Moved model to GPU")
                    except Exception as e:
                        logger.warning(f"Could not move model to GPU, using CPU: {e}")
                
                if model is None:
                    logger.error("Model loading failed")
                    cls._model_load_failed = True
                    cls._model = None
                    return None
                
                cls._model = model
                cls._model_loaded = True
                
                # Update embedding dimension in index manager
                dim = cls._model.get_sentence_embedding_dimension()
                FAISSIndexManager.EMBEDDING_DIM = dim
                logger.info(f"Embedding model ready: dim={dim}, device={cls._model.device}")
                
            except Exception as e:
                logger.error(f"Failed to load embedding model: {e}")
                import traceback
                logger.error(f"Traceback: {traceback.format_exc()}")
                cls._model = None
        return cls._model
    
    @classmethod
    def get_index_manager(cls) -> FAISSIndexManager:
        """Get or create FAISS index manager"""
        if cls._index_manager is None:
            cls._index_manager = FAISSIndexManager()
        return cls._index_manager
    
    def __init__(self):
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
        self.top_k = settings.TOP_K_RESULTS
    
    def _is_binary_or_base64(self, text: str) -> bool:
        """
        Check if text appears to be binary data or base64-encoded content.
        These should not be sent to the embedding model as they can cause crashes.
        """
        if not text or len(text) < 100:
            return False
        
        # Check for base64 image data URLs
        if 'data:image' in text or 'data:application' in text:
            return True
        
        # Check for very long base64-like strings (high ratio of alphanumeric + /+=)
        # Base64 uses A-Za-z0-9+/= characters
        sample = text[:1000] if len(text) > 1000 else text
        base64_chars = sum(1 for c in sample if c.isalnum() or c in '+/=')
        
        # If more than 90% of characters are base64-compatible and no spaces/newlines,
        # it's likely binary/base64 data
        if base64_chars / len(sample) > 0.9 and sample.count(' ') < 10 and sample.count('\n') < 5:
            # Additional check: real text has more variety in character distribution
            # Base64 tends to have very uniform distribution
            if len(set(sample)) < 70:  # Base64 only uses ~65 unique chars
                return True
        
        return False
    
    def _sanitize_text_for_embedding(self, text: str) -> str:
        """
        Remove or replace content that shouldn't be embedded (base64, binary, etc.)
        """
        import re
        
        # Remove base64 data URLs (images, files, etc.)
        # Pattern: data:mime/type;base64,<base64_data>
        text = re.sub(r'data:[^;]+;base64,[A-Za-z0-9+/=]+', '[IMAGE_DATA_REMOVED]', text)
        
        # Remove very long alphanumeric strings that look like base64 or hashes (> 200 chars)
        text = re.sub(r'[A-Za-z0-9+/=]{200,}', '[BINARY_DATA_REMOVED]', text)
        
        return text
    
    def embed_text(self, text: str) -> Optional[np.ndarray]:
        """Generate embedding for text using GPU"""
        model = self.get_model()
        if model is None:
            logger.warning("[EMBED] Model not loaded!")
            return None
        
        # Sanitize text to remove binary/base64 content that could crash the model
        if self._is_binary_or_base64(text):
            logger.warning(f"[EMBED] Skipping binary/base64 content (len={len(text)})")
            return None
        
        # Also sanitize inline base64 within otherwise normal text
        clean_text = self._sanitize_text_for_embedding(text)
        
        if not clean_text or len(clean_text.strip()) < 3:
            logger.warning("[EMBED] Text too short after sanitization")
            return None
        
        embedding = model.encode(clean_text, convert_to_numpy=True, normalize_embeddings=True)
        
        # Debug logging
        logger.debug(f"[EMBED] text='{clean_text[:50]}...', shape={embedding.shape}, norm={np.linalg.norm(embedding):.4f}, first5={embedding[:5].tolist()}")
        
        return embedding
    
    async def get_embedding(self, text: str) -> List[float]:
        """
        Async wrapper for embed_text that returns a list of floats.
        Used by the OpenAI-compatible /v1/embeddings endpoint.
        """
        import asyncio
        
        loop = asyncio.get_event_loop()
        embedding = await loop.run_in_executor(None, self.embed_text, text)
        
        if embedding is None:
            raise RuntimeError("Embedding model not available")
        
        return embedding.tolist()
    
    def embed_texts(self, texts: List[str], batch_size: int = 64) -> Optional[np.ndarray]:
        """Generate embeddings for multiple texts with batching"""
        model = self.get_model()
        if model is None:
            return None
        
        # Sanitize all texts - filter out binary/base64 content
        sanitized_texts = []
        valid_indices = []
        
        for i, text in enumerate(texts):
            if self._is_binary_or_base64(text):
                logger.warning(f"[EMBED_BATCH] Skipping binary/base64 content at index {i} (len={len(text)})")
                continue
            
            clean_text = self._sanitize_text_for_embedding(text)
            if clean_text and len(clean_text.strip()) >= 3:
                sanitized_texts.append(clean_text)
                valid_indices.append(i)
            else:
                logger.warning(f"[EMBED_BATCH] Text too short after sanitization at index {i}")
        
        if not sanitized_texts:
            logger.warning("[EMBED_BATCH] No valid texts to embed after sanitization")
            return None
        
        # Generate embeddings for sanitized texts
        embeddings = model.encode(
            sanitized_texts,
            convert_to_numpy=True,
            normalize_embeddings=True,
            batch_size=batch_size,
            show_progress_bar=len(sanitized_texts) > 100
        )
        
        # If we filtered some texts, we need to return None for those positions
        # to maintain alignment with the original texts list
        if len(valid_indices) < len(texts):
            import numpy as np
            full_embeddings = np.zeros((len(texts), embeddings.shape[1]), dtype=np.float32)
            for new_idx, orig_idx in enumerate(valid_indices):
                full_embeddings[orig_idx] = embeddings[new_idx]
            # Mark filtered positions with NaN so caller knows they're invalid
            for i in range(len(texts)):
                if i not in valid_indices:
                    full_embeddings[i] = np.nan
            return full_embeddings
        
        return embeddings
    
    def chunk_text(self, text: str) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks"""
        import re
        text = re.sub(r'\s+', ' ', text).strip()
        
        chunks = []
        words = text.split()
        
        if len(words) <= self.chunk_size:
            return [{"content": text, "index": 0}]
        
        start = 0
        chunk_index = 0
        
        while start < len(words):
            end = min(start + self.chunk_size, len(words))
            chunk_words = words[start:end]
            chunk_text = ' '.join(chunk_words)
            
            chunks.append({
                "content": chunk_text,
                "index": chunk_index,
            })
            
            chunk_index += 1
            start = end - self.chunk_overlap
            
            if start >= len(words) - self.chunk_overlap:
                break
        
        return chunks
    
    async def process_document(
        self,
        db: AsyncSession,
        document: Document,
        text_content: str,
        build_knowledge_graph: bool = True,
    ) -> int:
        """Process a document: chunk, embed, store, and build knowledge graph edges"""
        
        # Delete existing chunks
        await db.execute(
            delete(DocumentChunk).where(DocumentChunk.document_id == document.id)
        )
        
        # Chunk the text
        chunks = self.chunk_text(text_content)
        
        if not chunks:
            return 0
        
        # Generate embeddings (run in thread pool)
        chunk_texts = [c["content"] for c in chunks]
        embeddings = await asyncio.get_event_loop().run_in_executor(
            _executor,
            self.embed_texts,
            chunk_texts
        )
        
        # Store chunks in database
        new_chunks = []
        chunk_ids = []
        for i, chunk in enumerate(chunks):
            embedding_bytes = None
            if embeddings is not None:
                embedding_bytes = embeddings[i].astype(np.float32).tobytes()
            
            db_chunk = DocumentChunk(
                document_id=document.id,
                content=chunk["content"],
                chunk_index=chunk["index"],
                embedding=embedding_bytes,
                chunk_metadata={"word_count": len(chunk["content"].split())},
            )
            db.add(db_chunk)
            await db.flush()
            chunk_ids.append(str(db_chunk.id))
            new_chunks.append(db_chunk)
        
        # Update document status
        document.is_processed = True
        document.chunk_count = len(chunks)
        
        await db.flush()
        
        # Update FAISS index if document belongs to a knowledge store
        if document.knowledge_store_id and embeddings is not None:
            await self._update_knowledge_store_index(
                db, document.knowledge_store_id
            )
        
        # NC-0.8.0.2: Build knowledge graph edges for temporal validity
        if build_knowledge_graph and document.knowledge_store_id:
            try:
                await self._build_knowledge_graph_edges(db, new_chunks, document.knowledge_store_id)
            except Exception as e:
                logger.warning(f"Knowledge graph edge creation failed (non-fatal): {e}")
        
        return len(chunks)
    
    async def _build_knowledge_graph_edges(
        self,
        db: AsyncSession,
        new_chunks: List[DocumentChunk],
        knowledge_store_id: str,
    ) -> None:
        """
        Build knowledge graph edges by comparing new chunks against existing ones.
        
        NC-0.8.0.2: Detects SUPERSEDES, UPDATES, CONTRADICTS relationships
        """
        from app.services.knowledge_graph import get_knowledge_graph_service
        
        kg_service = get_knowledge_graph_service()
        
        # Get existing chunks from same knowledge store (excluding the new ones)
        new_chunk_ids = {str(c.id) for c in new_chunks}
        
        result = await db.execute(
            select(DocumentChunk)
            .join(Document, DocumentChunk.document_id == Document.id)
            .where(Document.knowledge_store_id == knowledge_store_id)
            .where(Document.is_processed == True)
            .where(DocumentChunk.id.notin_(new_chunk_ids))
            .limit(100)  # Limit for performance - could do batching for large KBs
        )
        existing_chunks = result.scalars().all()
        
        if not existing_chunks:
            logger.debug(f"No existing chunks to compare against in KB {knowledge_store_id}")
            return
        
        logger.info(f"[KNOWLEDGE_GRAPH] Analyzing {len(new_chunks)} new chunks against {len(existing_chunks)} existing")
        
        total_edges = 0
        for new_chunk in new_chunks:
            # Analyze this chunk against existing (with LLM if available, otherwise heuristics)
            edges = await kg_service.analyze_new_chunk(
                db, 
                new_chunk, 
                existing_chunks,
                use_llm=False  # Start with heuristics for speed; can enable LLM later
            )
            total_edges += len(edges)
            
            # Also extract entities for future relationship detection
            await kg_service.extract_entities(db, new_chunk, use_llm=False)
        
        if total_edges > 0:
            logger.info(f"[KNOWLEDGE_GRAPH] Created {total_edges} edges for new chunks")
        
        await db.flush()
    
    async def delete_document(
        self,
        db: AsyncSession,
        document_id: str,
    ) -> None:
        """Delete a document and its chunks, update FAISS indexes"""
        
        # Get document to check knowledge store
        result = await db.execute(
            select(Document).where(Document.id == document_id)
        )
        document = result.scalar_one_or_none()
        
        if not document:
            return  # Document already deleted
        
        knowledge_store_id = document.knowledge_store_id
        
        # Delete chunks first (explicit delete for reliability)
        await db.execute(
            delete(DocumentChunk).where(DocumentChunk.document_id == document_id)
        )
        
        # Delete document using ORM to trigger cascades
        await db.delete(document)
        
        # Rebuild FAISS index if document was in a knowledge store
        if knowledge_store_id:
            try:
                await self._update_knowledge_store_index(db, knowledge_store_id)
            except Exception as e:
                logger.warning(f"Failed to update FAISS index after document deletion: {e}")
    
    async def _update_knowledge_store_index(
        self,
        db: AsyncSession,
        knowledge_store_id: str,
    ) -> None:
        """Rebuild FAISS index for a knowledge store"""
        
        logger.info(f"[FAISS_BUILD] Starting index rebuild for knowledge store: {knowledge_store_id}")
        
        # Get all chunks for this knowledge store
        result = await db.execute(
            select(DocumentChunk, Document)
            .join(Document, DocumentChunk.document_id == Document.id)
            .where(Document.knowledge_store_id == knowledge_store_id)
            .where(Document.is_processed == True)
            .where(DocumentChunk.embedding.isnot(None))
        )
        rows = result.all()
        
        if not rows:
            logger.warning(f"[FAISS_BUILD] No chunks found for knowledge store {knowledge_store_id}")
            self.get_index_manager().delete_index(knowledge_store_id)
            return
        
        logger.info(f"[FAISS_BUILD] Found {len(rows)} chunks to index")
        
        # Collect embeddings and IDs
        embeddings = []
        chunk_ids = []
        
        for chunk, document in rows:
            embedding = np.frombuffer(chunk.embedding, dtype=np.float32)
            embeddings.append(embedding)
            chunk_ids.append(str(chunk.id))
        
        embeddings_array = np.vstack(embeddings)
        
        logger.info(f"[FAISS_BUILD] Embedding array shape: {embeddings_array.shape}, first embedding norm: {np.linalg.norm(embeddings_array[0]):.4f}")
        
        # Build index in thread pool
        await asyncio.get_event_loop().run_in_executor(
            _executor,
            self.get_index_manager().build_index,
            knowledge_store_id,
            embeddings_array,
            chunk_ids
        )
        
        logger.info(f"[FAISS_BUILD] Index build complete for {knowledge_store_id}")
    
    async def search(
        self,
        db: AsyncSession,
        user: User,
        query: str,
        document_ids: Optional[List[str]] = None,
        top_k: Optional[int] = None,
    ) -> List[Dict[str, Any]]:
        """
        Search for relevant chunks using FAISS.
        
        Args:
            document_ids: If provided, search only in these specific documents.
                         If None, search user's unitemized documents (not in any knowledge store).
        
        Note: Documents in knowledge stores are only searched via search_knowledge_stores(),
        which requires explicit knowledge store IDs.
        """
        top_k = top_k or self.top_k
        
        # Check if debug logging is enabled
        debug_enabled = await _is_debug_rag_enabled(db)
        
        if debug_enabled:
            _log_rag_debug(
                "Search started",
                user_id=user.id,
                query=query[:100] + "..." if len(query) > 100 else query,
                document_ids=document_ids,
                top_k=top_k
            )
        
        # Generate query embedding
        query_embedding = await asyncio.get_event_loop().run_in_executor(
            _executor, self.embed_text, query
        )
        
        if query_embedding is None:
            if debug_enabled:
                _log_rag_debug("Falling back to keyword search (embedding failed)")
            return await self._keyword_search(db, user, query, document_ids, top_k)
        
        # If document_ids specified, do direct search
        if document_ids:
            results = await self._direct_search(db, query_embedding, document_ids, top_k)
            if debug_enabled:
                _log_rag_debug(
                    "Direct search completed",
                    results_count=len(results),
                    top_scores=[r.get("similarity", 0) for r in results[:3]]
                )
            return results
        
        # Otherwise search user's documents
        results = await self._user_document_search(db, user, query_embedding, top_k)
        if debug_enabled:
            _log_rag_debug(
                "User document search completed",
                results_count=len(results),
                top_scores=[r.get("similarity", 0) for r in results[:3]]
            )
        return results
    
    async def _direct_search(
        self,
        db: AsyncSession,
        query_embedding: np.ndarray,
        document_ids: List[str],
        top_k: int,
    ) -> List[Dict[str, Any]]:
        """Direct search in specific documents (no FAISS index)"""
        query_builder = (
            select(DocumentChunk, Document)
            .join(Document, DocumentChunk.document_id == Document.id)
            .where(Document.id.in_(document_ids))
            .where(DocumentChunk.embedding.isnot(None))
        )
        
        result = await db.execute(query_builder)
        rows = result.all()
        
        if not rows:
            logger.debug(f"[DIRECT_SEARCH] No chunks found for documents: {document_ids}")
            return []
        
        # Ensure query is normalized
        query_norm = np.linalg.norm(query_embedding)
        if query_norm > 0:
            query_normalized = query_embedding / query_norm
        else:
            logger.warning("[DIRECT_SEARCH] Query embedding has zero norm!")
            return []
        
        logger.debug(f"[DIRECT_SEARCH] Found {len(rows)} chunks, query_norm_before={query_norm:.4f}, query_dim={len(query_embedding)}")
        
        # Calculate similarities
        results = []
        for chunk, document in rows:
            chunk_embedding = np.frombuffer(chunk.embedding, dtype=np.float32)
            
            # Check dimensions match
            if len(chunk_embedding) != len(query_normalized):
                logger.warning(f"[DIRECT_SEARCH] Dimension mismatch! chunk_dim={len(chunk_embedding)}, query_dim={len(query_normalized)}")
                continue
            
            # Normalize chunk embedding
            chunk_norm = np.linalg.norm(chunk_embedding)
            if chunk_norm > 0:
                chunk_normalized = chunk_embedding / chunk_norm
            else:
                logger.warning(f"[DIRECT_SEARCH] Chunk has zero norm: {chunk.id}")
                continue
            
            # Cosine similarity
            similarity = float(np.dot(query_normalized, chunk_normalized))
            
            if len(results) < 3:  # Log first 3
                logger.debug(f"[DIRECT_SEARCH] Chunk '{chunk.content[:50]}...' norm={chunk_norm:.4f}, similarity={similarity:.4f}")
            
            results.append({
                "chunk_id": str(chunk.id),
                "document_id": str(document.id),
                "document_name": document.name,
                "content": chunk.content,
                "chunk_index": chunk.chunk_index,
                "similarity": similarity,
                "metadata": chunk.chunk_metadata,
            })
        
        results.sort(key=lambda x: x["similarity"], reverse=True)
        logger.debug(f"[DIRECT_SEARCH] Top similarities: {[round(r['similarity'], 4) for r in results[:5]]}")
        return results[:top_k]
    
    async def _user_document_search(
        self,
        db: AsyncSession,
        user: User,
        query_embedding: np.ndarray,
        top_k: int,
    ) -> List[Dict[str, Any]]:
        """
        Search user's unitemized documents (not in any knowledge store).
        
        Documents in knowledge stores are only searched when:
        - A custom GPT with that knowledge store is selected
        - The user explicitly searches a knowledge store
        
        This prevents searching subscribed GPT knowledge stores when no GPT is selected.
        """
        # Get user's documents that are NOT in any knowledge store
        result = await db.execute(
            select(Document)
            .where(Document.owner_id == user.id)
            .where(Document.is_processed == True)
            .where(Document.knowledge_store_id.is_(None))  # Only unitemized documents
        )
        documents = result.scalars().all()
        
        if not documents:
            logger.debug(f"[USER_DOC_SEARCH] No unitemized documents found for user {user.id}")
            return []
        
        logger.debug(f"[USER_DOC_SEARCH] Found {len(documents)} unitemized documents for user {user.id}")
        
        doc_ids = [str(d.id) for d in documents]
        return await self._direct_search(db, query_embedding, doc_ids, top_k)
    
    async def search_knowledge_stores(
        self,
        db: AsyncSession,
        user: User,
        query: str,
        knowledge_store_ids: List[str],
        top_k: Optional[int] = None,
        bypass_access_check: bool = False,
    ) -> List[Dict[str, Any]]:
        """
        Search across knowledge stores using FAISS indexes with hybrid search.
        
        Uses hybrid search combining:
        1. Semantic search (FAISS cosine similarity)
        2. Identifier search (exact substring matching for codes/statutes)
        
        Args:
            bypass_access_check: If True, skip ownership/permission checks.
                                Used when accessing KB through a public GPT.
        """
        if not knowledge_store_ids:
            return []
        
        top_k = top_k or self.top_k
        
        # Extract identifiers for hybrid search
        identifiers = self._extract_identifiers(query)
        
        # Check if debug logging is enabled
        debug_enabled = await _is_debug_rag_enabled(db)
        
        if debug_enabled:
            # Get knowledge store names for logging
            kb_result = await db.execute(
                select(KnowledgeStore).where(KnowledgeStore.id.in_(knowledge_store_ids))
            )
            kb_names = {str(kb.id): kb.name for kb in kb_result.scalars().all()}
            
            _log_rag_debug(
                "Knowledge store search started",
                user_id=user.id,
                query=query[:100] + "..." if len(query) > 100 else query,
                knowledge_stores=[kb_names.get(kid, kid) for kid in knowledge_store_ids],
                top_k=top_k,
                bypass_access_check=bypass_access_check,
                identifiers=identifiers if identifiers else None
            )
        
        # Verify access (unless bypassed for assistant access)
        if bypass_access_check:
            accessible_stores = knowledge_store_ids
        else:
            accessible_stores = await self._get_accessible_store_ids(db, user, knowledge_store_ids)
        
        if not accessible_stores:
            if debug_enabled:
                _log_rag_debug("No accessible knowledge stores found")
            return []
        
        if debug_enabled and len(accessible_stores) != len(knowledge_store_ids):
            _log_rag_debug(
                "Access filtered stores",
                requested=len(knowledge_store_ids),
                accessible=len(accessible_stores)
            )
        
        # Run identifier search if identifiers were found
        identifier_results = []
        if identifiers:
            identifier_results = await self._identifier_search(
                db, identifiers, knowledge_store_ids=accessible_stores, top_k=top_k
            )
            if debug_enabled and identifier_results:
                _log_rag_debug(
                    "Identifier search results",
                    count=len(identifier_results),
                    identifiers=identifiers
                )
        
        # Generate query embedding
        query_embedding = await asyncio.get_event_loop().run_in_executor(
            _executor, self.embed_text, query
        )
        
        if query_embedding is None:
            if debug_enabled:
                _log_rag_debug("Falling back to keyword search (embedding failed)")
            # If we have identifier results, return those
            if identifier_results:
                return identifier_results[:top_k]
            # Fallback to keyword search
            doc_ids = await self._get_store_document_ids(db, accessible_stores)
            return await self._keyword_search_documents(db, doc_ids, query, top_k)
        
        # Search each knowledge store's FAISS index
        all_results = []
        index_manager = self.get_index_manager()
        
        logger.info(f"[KB_SEARCH] Searching {len(accessible_stores)} knowledge stores: {accessible_stores}")
        
        for store_id in accessible_stores:
            logger.info(f"[KB_SEARCH] Searching store: {store_id} (type: {type(store_id).__name__})")
            results = await asyncio.get_event_loop().run_in_executor(
                _executor,
                index_manager.search,
                store_id,
                query_embedding,
                top_k
            )
            logger.info(f"[KB_SEARCH] Store {store_id} returned {len(results)} results")
            
            if debug_enabled and results:
                _log_rag_debug(
                    f"FAISS index search for store",
                    store_id=store_id,
                    results_count=len(results),
                    top_score=results[0][1] if results else 0
                )
            
            for chunk_id, score in results:
                all_results.append((chunk_id, score, store_id))
        
        # Get semantic results
        if all_results:
            all_results.sort(key=lambda x: x[1], reverse=True)
            top_results = all_results[:top_k]
            chunk_ids = [r[0] for r in top_results]
            semantic_results = await self._get_chunks_by_ids(db, chunk_ids, {r[0]: r[1] for r in top_results})
        else:
            semantic_results = []
        
        # Merge semantic and identifier results
        if identifiers and identifier_results:
            final_results = self._merge_search_results(
                semantic_results,
                identifier_results,
                semantic_threshold=0.7,
                identifier_boost=0.15,
                top_k=top_k
            )
        else:
            final_results = semantic_results
        
        if not final_results:
            if debug_enabled:
                _log_rag_debug("No FAISS results, falling back to direct search")
            # Fallback to direct search
            doc_ids = await self._get_store_document_ids(db, accessible_stores)
            return await self._direct_search(db, query_embedding, doc_ids, top_k)
        
        # NC-0.8.0.2: Apply knowledge graph temporal validity filtering
        try:
            from app.services.knowledge_graph import get_knowledge_graph_service
            kg_service = get_knowledge_graph_service()
            
            # Classify query intent for temporal awareness
            temporal_intent = kg_service.classify_temporal_intent(query)
            
            if debug_enabled:
                _log_rag_debug(
                    "Temporal intent classification",
                    is_time_sensitive=temporal_intent["is_time_sensitive"],
                    intent_type=temporal_intent["intent_type"],
                    confidence=temporal_intent["confidence"]
                )
            
            # Apply validity filtering
            final_results = await kg_service.filter_by_validity(
                db, final_results, query, temporal_intent
            )
            
            if debug_enabled:
                _log_rag_debug(
                    "Knowledge graph filtering complete",
                    results_after_filter=len(final_results),
                    any_superseded=any(r.get("validity", {}).get("superseded_by") for r in final_results)
                )
        except Exception as e:
            logger.warning(f"Knowledge graph filtering failed (continuing without): {e}")
        
        if debug_enabled:
            _log_rag_debug(
                "Knowledge store search completed",
                total_results=len(final_results),
                top_scores=[round(r.get("similarity", 0), 4) for r in final_results[:5]]
            )
        
        if debug_enabled and final_results:
            _log_rag_debug(
                "Retrieved context chunks",
                chunks=[{
                    "doc": r.get("document_name", "?"),
                    "score": round(r.get("similarity", 0), 4),
                    "preview": r.get("content", "")[:80] + "..."
                } for r in final_results[:3]]
            )
        
        return final_results
    
    async def _enhance_query_with_context(
        self,
        db: AsyncSession,
        query: str,
        chat_id: Optional[str] = None,
        max_context_messages: int = 3,
    ) -> str:
        """
        Enhance a vague query with context from recent conversation.
        
        When user asks follow-up questions like "Where did it come from?" or 
        "Tell me more about that", this extracts keywords from recent messages
        to make the RAG search more effective.
        
        Args:
            query: The current user query
            chat_id: Optional chat ID to get context from
            max_context_messages: Number of recent messages to consider
            
        Returns:
            Enhanced query string combining original query with context keywords
        """
        # Only enhance short/vague queries (likely follow-ups)
        if len(query.split()) > 10 or not chat_id:
            return query
        
        # Common follow-up indicators
        follow_up_words = {
            'it', 'that', 'this', 'they', 'them', 'those', 'these',
            'there', 'where', 'when', 'how', 'why', 'what', 'who',
            'more', 'also', 'too', 'another', 'other', 'same',
        }
        
        query_words = set(query.lower().split())
        is_likely_followup = len(query_words & follow_up_words) > 0 or len(query_words) <= 5
        
        if not is_likely_followup:
            return query
        
        try:
            from app.models.models import Message
            
            # Get recent messages from this chat
            result = await db.execute(
                select(Message)
                .where(Message.chat_id == chat_id)
                .order_by(Message.created_at.desc())
                .limit(max_context_messages * 2)  # Get more to filter
            )
            messages = result.scalars().all()
            
            if not messages:
                return query
            
            # Extract keywords from recent messages (both user and assistant)
            # Focus on nouns and proper nouns - words that are likely topics
            import re
            context_keywords = set()
            
            # Skip common words that don't add search value
            stop_words = {
                'the', 'a', 'an', 'is', 'are', 'was', 'were', 'be', 'been',
                'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will',
                'would', 'could', 'should', 'may', 'might', 'must', 'shall',
                'can', 'need', 'dare', 'ought', 'used', 'to', 'of', 'in',
                'for', 'on', 'with', 'at', 'by', 'from', 'as', 'into', 'through',
                'during', 'before', 'after', 'above', 'below', 'between', 'under',
                'again', 'further', 'then', 'once', 'here', 'there', 'when',
                'where', 'why', 'how', 'all', 'each', 'few', 'more', 'most',
                'other', 'some', 'such', 'no', 'nor', 'not', 'only', 'own',
                'same', 'so', 'than', 'too', 'very', 'just', 'and', 'but',
                'if', 'or', 'because', 'until', 'while', 'about', 'against',
                'i', 'me', 'my', 'myself', 'we', 'our', 'ours', 'ourselves',
                'you', 'your', 'yours', 'yourself', 'yourselves', 'he', 'him',
                'his', 'himself', 'she', 'her', 'hers', 'herself', 'it', 'its',
                'itself', 'they', 'them', 'their', 'theirs', 'themselves',
                'what', 'which', 'who', 'whom', 'this', 'that', 'these', 'those',
                'am', 'tell', 'know', 'think', 'want', 'see', 'look', 'make',
                'go', 'get', 'come', 'take', 'give', 'say', 'said', 'like',
            }
            
            for msg in messages[:max_context_messages]:
                if not msg.content:
                    continue
                    
                # Extract words, focusing on capitalized words (likely proper nouns/topics)
                # and longer words (more likely to be meaningful)
                words = re.findall(r'\b[A-Za-z][A-Za-z0-9-]{2,}\b', msg.content)
                
                for word in words:
                    lower_word = word.lower()
                    # Include if: capitalized (proper noun), longer than 5 chars, or contains numbers
                    if (lower_word not in stop_words and 
                        (word[0].isupper() or len(word) > 5 or any(c.isdigit() for c in word))):
                        context_keywords.add(word)
            
            if not context_keywords:
                return query
            
            # Take top keywords (prefer proper nouns/capitalized, then by length)
            sorted_keywords = sorted(
                context_keywords,
                key=lambda w: (not w[0].isupper(), -len(w))
            )[:5]
            
            # Combine query with context keywords
            enhanced_query = f"{query} {' '.join(sorted_keywords)}"
            logger.info(f"[RAG_CONTEXT] Enhanced query: '{query}' -> '{enhanced_query}'")
            
            return enhanced_query
            
        except Exception as e:
            logger.warning(f"[RAG_CONTEXT] Failed to enhance query: {e}")
            return query
    
    def _normalize_text_for_matching(self, text: str) -> str:
        """
        Normalize text for identifier matching by converting Unicode variants
        and common word patterns to their ASCII equivalents.
        
        This handles common cases where:
        - Documents contain typographic characters (em dashes, curly quotes)
        - Users search with regular keyboard characters
        - Users write ranges as words ("1 to 9") instead of dashes ("1-9")
        
        Normalizations:
        - All dash variants → hyphen-minus (-)
        - All quote variants → straight quotes (' and ")
        - Range words "to", "through", "thru" between tokens with digits → hyphen
        """
        import re
        
        # Dash/hyphen variants → regular hyphen
        # U+002D HYPHEN-MINUS, U+2010 HYPHEN, U+2011 NON-BREAKING HYPHEN,
        # U+2012 FIGURE DASH, U+2013 EN DASH, U+2014 EM DASH,
        # U+2015 HORIZONTAL BAR, U+2212 MINUS SIGN
        dash_pattern = r'[\u2010\u2011\u2012\u2013\u2014\u2015\u2212]'
        text = re.sub(dash_pattern, '-', text)
        
        # Quote variants → straight quotes (for identifiers that might have quotes)
        # Single quotes: U+2018, U+2019, U+201A, U+201B
        text = re.sub(r'[\u2018\u2019\u201a\u201b]', "'", text)
        # Double quotes: U+201C, U+201D, U+201E, U+201F
        text = re.sub(r'[\u201c\u201d\u201e\u201f]', '"', text)
        
        # Convert range words between tokens to hyphens
        # "1 to 9" → "1-9", "Rule 1 through 5" → "Rule 1-5"
        # Requirements:
        # - At least one side must contain a digit (to avoid "go to the" → "go-the")
        # - Tokens can be alphanumeric (letters and/or numbers)
        # Pattern matches: (token_with_digit) to (any_token) OR (any_token) to (token_with_digit)
        
        def replace_range(match):
            left, right = match.group(1), match.group(2)
            # Only replace if at least one side has a digit
            if any(c.isdigit() for c in left) or any(c.isdigit() for c in right):
                return f"{left}-{right}"
            return match.group(0)  # Return unchanged
        
        range_pattern = r'(\b[A-Za-z0-9]+)\s+(?:to|through|thru)\s+([A-Za-z0-9]+\b)'
        text = re.sub(range_pattern, replace_range, text, flags=re.IGNORECASE)
        
        return text
    
    def _extract_identifiers(self, query: str) -> List[str]:
        """
        Extract specific identifiers from a query that should be searched exactly.
        
        These patterns typically fail in pure semantic search because they're
        alphanumeric codes that don't have strong semantic meaning on their own.
        
        Examples matched:
        - Legal: "1S-2.053", "Rule 12-345.67", "FAC 61G15-30"
        - Statutes: "§ 119.07", "Section 768.28"  
        - Codes: "ICD-10-CM", "CPT-99213"
        - Patents/Cases: "US-2024-123456", "Case No. 2024-CF-001234"
        - Product codes: "ABC-123-XYZ", "SKU-78910"
        - UUIDs: partial matches like "a1b2c3d4"
        
        Returns list of identifier strings to search for exactly.
        """
        import re
        
        # Normalize dashes/hyphens in query before extraction
        # Unicode dash variants that should all be treated as hyphens:
        # - U+002D HYPHEN-MINUS (regular hyphen)
        # - U+2010 HYPHEN
        # - U+2011 NON-BREAKING HYPHEN
        # - U+2012 FIGURE DASH
        # - U+2013 EN DASH
        # - U+2014 EM DASH
        # - U+2015 HORIZONTAL BAR
        # - U+2212 MINUS SIGN
        normalized_query = self._normalize_text_for_matching(query)
        
        identifiers = []
        
        # Pattern 1: Alphanumeric with dots/dashes (statutes, codes, rules)
        # Matches: "1S-2.053", "61G15-30.001", "768.28", "119.07(1)"
        pattern_statute = r'\b\d+[A-Z]?[-.][\dA-Z]+(?:[-.]\d+)*(?:\([a-z0-9]+\))*\b'
        identifiers.extend(re.findall(pattern_statute, normalized_query, re.IGNORECASE))
        
        # Pattern 2: Letter-prefixed codes with numbers
        # Matches: "FAC 61G15", "Rule 12-345", "ICD-10", "CPT-99213"
        pattern_codes = r'\b[A-Z]{2,}[-\s]?\d+[A-Z]?(?:[-.][\dA-Z]+)*\b'
        identifiers.extend(re.findall(pattern_codes, normalized_query, re.IGNORECASE))
        
        # Pattern 3: Section/paragraph references
        # Matches: "§ 119.07", "§119.07", "Section 768.28"
        pattern_section = r'(?:§\s*|[Ss]ection\s+)(\d+(?:\.\d+)*(?:\([a-z0-9]+\))*)'
        for match in re.finditer(pattern_section, normalized_query):
            identifiers.append(match.group(1))
        
        # Pattern 4: Case/docket numbers
        # Matches: "2024-CF-001234", "Case 23-cv-12345"
        pattern_case = r'\b\d{2,4}[-][A-Z]{1,3}[-]\d{4,}\b'
        identifiers.extend(re.findall(pattern_case, normalized_query, re.IGNORECASE))
        
        # Pattern 5: General alphanumeric identifiers with at least one letter and number
        # and containing a separator (dash or dot)
        # Matches: "ABC-123", "A1B2-C3D4"
        pattern_general = r'\b(?=[A-Z0-9]*[A-Z])(?=[A-Z0-9]*\d)[A-Z0-9]+[-./][A-Z0-9]+(?:[-./][A-Z0-9]+)*\b'
        identifiers.extend(re.findall(pattern_general, normalized_query, re.IGNORECASE))
        
        # Deduplicate while preserving order, normalize to lowercase for matching
        seen = set()
        unique_identifiers = []
        for ident in identifiers:
            # Clean up and normalize
            ident_clean = ident.strip()
            ident_lower = ident_clean.lower()
            if ident_lower not in seen and len(ident_clean) >= 3:
                seen.add(ident_lower)
                unique_identifiers.append(ident_clean)
        
        if unique_identifiers:
            logger.info(f"[RAG_HYBRID] Extracted identifiers from query: {unique_identifiers}")
        
        return unique_identifiers
    
    async def _identifier_search(
        self,
        db: AsyncSession,
        identifiers: List[str],
        knowledge_store_ids: Optional[List[str]] = None,
        top_k: int = 10,
    ) -> List[Dict[str, Any]]:
        """
        Search for chunks containing specific identifiers using exact substring matching.
        
        This complements semantic search by finding chunks that contain specific
        alphanumeric codes/identifiers that semantic search may miss due to
        low semantic similarity.
        
        Args:
            identifiers: List of identifier strings to search for
            knowledge_store_ids: Optional list of knowledge store IDs to search within
            top_k: Maximum results to return
            
        Returns:
            List of matching chunks with similarity scores based on match quality
        """
        if not identifiers:
            return []
        
        # Build query for chunks
        query_builder = (
            select(DocumentChunk, Document)
            .join(Document, DocumentChunk.document_id == Document.id)
            .where(Document.is_processed == True)
        )
        
        # Filter by knowledge stores if specified
        if knowledge_store_ids:
            query_builder = query_builder.where(
                Document.knowledge_store_id.in_(knowledge_store_ids)
            )
        
        result = await db.execute(query_builder)
        rows = result.all()
        
        if not rows:
            return []
        
        results = []
        for chunk, document in rows:
            # Normalize content to handle Unicode dash/quote variants
            # This allows matching "1-9" in query against "1—9" (em dash) in document
            content_normalized = self._normalize_text_for_matching(chunk.content).lower()
            
            # Check each identifier
            matches = []
            for ident in identifiers:
                ident_lower = ident.lower()
                if ident_lower in content_normalized:
                    # Calculate match quality
                    # Exact match at word boundary scores higher
                    import re
                    if re.search(r'\b' + re.escape(ident_lower) + r'\b', content_normalized):
                        matches.append((ident, 1.0))  # Perfect word boundary match
                    else:
                        matches.append((ident, 0.8))  # Substring match
            
            if matches:
                # Score based on: number of identifiers matched and match quality
                match_score = sum(score for _, score in matches) / len(identifiers)
                # Boost for multiple matches
                if len(matches) > 1:
                    match_score = min(1.0, match_score * 1.2)
                
                results.append({
                    "chunk_id": str(chunk.id),
                    "document_id": str(document.id),
                    "document_name": document.name,
                    "knowledge_store_id": str(document.knowledge_store_id) if document.knowledge_store_id else None,
                    "content": chunk.content,
                    "chunk_index": chunk.chunk_index,
                    "similarity": match_score,
                    "metadata": chunk.chunk_metadata,
                    "_match_type": "identifier",  # Internal marker
                    "_matched_identifiers": [ident for ident, _ in matches],
                })
        
        results.sort(key=lambda x: x["similarity"], reverse=True)
        
        if results:
            logger.info(f"[RAG_HYBRID] Identifier search found {len(results)} matches for: {identifiers}")
        
        return results[:top_k]
    
    def _merge_search_results(
        self,
        semantic_results: List[Dict[str, Any]],
        identifier_results: List[Dict[str, Any]],
        semantic_threshold: float = 0.7,
        identifier_boost: float = 0.15,
        top_k: int = 10,
    ) -> List[Dict[str, Any]]:
        """
        Merge semantic search results with identifier search results.
        
        Strategy:
        1. Identifier matches get a score boost since they're exact matches
        2. If a chunk appears in both, use the higher (boosted) score
        3. Identifier matches below semantic threshold are still included
           (this is the key improvement - they won't be filtered out)
        
        Args:
            semantic_results: Results from FAISS semantic search
            identifier_results: Results from exact identifier matching
            semantic_threshold: The threshold used for semantic search
            identifier_boost: Boost added to identifier match scores
            top_k: Maximum results to return
        """
        # Create lookup by chunk_id
        merged = {}
        
        # Add semantic results first
        for result in semantic_results:
            chunk_id = result["chunk_id"]
            merged[chunk_id] = result.copy()
            merged[chunk_id]["_sources"] = ["semantic"]
        
        # Merge identifier results with boost
        for result in identifier_results:
            chunk_id = result["chunk_id"]
            boosted_score = min(1.0, result["similarity"] + identifier_boost)
            
            if chunk_id in merged:
                # Chunk found by both methods - use higher score
                existing_score = merged[chunk_id]["similarity"]
                if boosted_score > existing_score:
                    merged[chunk_id]["similarity"] = boosted_score
                    merged[chunk_id]["_match_type"] = "hybrid"
                merged[chunk_id]["_sources"].append("identifier")
                if "_matched_identifiers" in result:
                    merged[chunk_id]["_matched_identifiers"] = result["_matched_identifiers"]
            else:
                # Only found by identifier search - include it!
                result_copy = result.copy()
                result_copy["similarity"] = boosted_score
                result_copy["_sources"] = ["identifier"]
                merged[chunk_id] = result_copy
        
        # Sort by score and return top_k
        results = list(merged.values())
        results.sort(key=lambda x: x["similarity"], reverse=True)
        
        # Log merge results
        hybrid_count = sum(1 for r in results if r.get("_sources") and len(r["_sources"]) > 1)
        identifier_only = sum(1 for r in results if r.get("_sources") == ["identifier"])
        if identifier_only > 0 or hybrid_count > 0:
            logger.info(
                f"[RAG_HYBRID] Merged results: {len(results)} total, "
                f"{hybrid_count} hybrid (both methods), {identifier_only} identifier-only"
            )
        
        return results[:top_k]
    
    async def search_global_stores(
        self,
        db: AsyncSession,
        query: str,
        chat_id: Optional[str] = None,
    ) -> Tuple[List[Dict[str, Any]], List[str]]:
        """
        Search all global knowledge stores for relevant context.
        
        Uses hybrid search combining:
        1. Semantic search (FAISS cosine similarity)
        2. Identifier search (exact substring matching for codes/statutes)
        
        Args:
            query: The search query
            chat_id: Optional chat ID for context-aware query enhancement
        
        Returns:
            Tuple of (results, store_names) where results are filtered by 
            each store's global_min_score threshold.
            
        This function is called automatically on every chat message when
        the global_knowledge_store_enabled setting is true.
        """
        import time
        from app.models.models import SystemSetting
        
        total_start = time.time()
        
        # Enhance query with conversation context (for follow-up questions)
        original_query = query
        if chat_id:
            query = await self._enhance_query_with_context(db, query, chat_id)
        
        # Extract identifiers for hybrid search
        identifiers = self._extract_identifiers(original_query)
        
        # Check if global knowledge stores are enabled
        setting_start = time.time()
        result = await db.execute(
            select(SystemSetting).where(SystemSetting.key == "global_knowledge_store_enabled")
        )
        setting = result.scalar_one_or_none()
        setting_time = time.time() - setting_start
        
        # Log whether setting exists and its value
        if setting:
            logger.info(f"[GLOBAL_RAG] Setting 'global_knowledge_store_enabled' = '{setting.value}'")
            if setting.value.lower() != "true":
                logger.info("[GLOBAL_RAG] Global knowledge stores disabled by setting")
                return [], []
        else:
            logger.info("[GLOBAL_RAG] Setting 'global_knowledge_store_enabled' not found, defaulting to enabled")
        
        # Find all global knowledge stores
        stores_start = time.time()
        result = await db.execute(
            select(KnowledgeStore).where(KnowledgeStore.is_global == True)
        )
        global_stores = result.scalars().all()
        stores_time = time.time() - stores_start
        logger.debug(f"[GLOBAL_RAG_TIMING] Store query took {stores_time:.3f}s, found {len(global_stores)} stores")
        
        if not global_stores:
            logger.debug("[GLOBAL_RAG] No global knowledge stores configured")
            return [], []
        
        # NC-0.8.0.1: Filter stores by required keywords
        # Only include stores where keywords match (or keyword filtering is disabled)
        query_lower = query.lower()
        original_query_lower = original_query.lower()
        
        def store_matches_keywords(store) -> bool:
            """Check if store's required keywords match the query."""
            if not store.require_keywords_enabled:
                return True  # No keyword filter, always include
            
            keywords = store.required_keywords
            if not keywords or not isinstance(keywords, list) or len(keywords) == 0:
                return True  # No keywords defined, always include
            
            # Check if any keyword/phrase is present in the query
            for keyword in keywords:
                if not isinstance(keyword, str):
                    continue
                keyword_lower = keyword.lower().strip()
                if keyword_lower and (keyword_lower in query_lower or keyword_lower in original_query_lower):
                    logger.debug(f"[GLOBAL_RAG] Store '{store.name}' keyword matched: '{keyword}'")
                    return True
            
            return False
        
        relevant_stores = [s for s in global_stores if store_matches_keywords(s)]
        skipped_stores = len(global_stores) - len(relevant_stores)
        
        if skipped_stores > 0:
            logger.info(f"[GLOBAL_RAG] Skipped {skipped_stores} stores due to keyword filter (no match)")
        
        if not relevant_stores:
            logger.debug("[GLOBAL_RAG] No global stores matched keyword requirements")
            return [], []
        
        logger.info(f"[GLOBAL_RAG] Searching {len(relevant_stores)} global stores for query: {query[:50]}...")
        
        global_store_ids = [str(store.id) for store in relevant_stores]
        
        # Run identifier search if identifiers were found
        identifier_results = []
        if identifiers:
            ident_start = time.time()
            identifier_results = await self._identifier_search(
                db, identifiers, knowledge_store_ids=global_store_ids, top_k=10
            )
            ident_time = time.time() - ident_start
            logger.debug(f"[GLOBAL_RAG_TIMING] Identifier search took {ident_time:.3f}s, found {len(identifier_results)} matches")
        
        # Generate query embedding for semantic search
        embed_start = time.time()
        query_embedding = await asyncio.get_event_loop().run_in_executor(
            _executor, self.embed_text, query
        )
        embed_time = time.time() - embed_start
        logger.debug(f"[GLOBAL_RAG_TIMING] Embedding generation took {embed_time:.3f}s")
        
        if query_embedding is None:
            logger.warning("[GLOBAL_RAG] Failed to generate embedding for query")
            # If we have identifier results, return those
            if identifier_results:
                store_names_with_matches = list(set(
                    r.get("_store_name", "") for r in identifier_results 
                    if r.get("knowledge_store_id") in global_store_ids
                ))
                return identifier_results, [s.name for s in global_stores if str(s.id) in 
                    set(r.get("knowledge_store_id") for r in identifier_results)]
            return [], []
        
        all_semantic_results = []
        matched_store_names = []
        index_manager = self.get_index_manager()
        
        # Get the lowest min_score among all relevant stores (for initial filter)
        # We'll do per-store filtering after merging
        global_min_score = min((store.global_min_score or 0.7) for store in relevant_stores)
        
        for store in relevant_stores:
            store_start = time.time()
            store_id = str(store.id)
            min_score = store.global_min_score or 0.7
            max_results = store.global_max_results or 3
            
            # Search this store's FAISS index
            results = await asyncio.get_event_loop().run_in_executor(
                _executor,
                index_manager.search,
                store_id,
                query_embedding,
                max_results * 2  # Get extra to filter by score
            )
            store_search_time = time.time() - store_start
            
            # Filter by minimum score threshold
            filtered_results = [(chunk_id, score) for chunk_id, score in results if score >= min_score]
            
            logger.debug(f"[GLOBAL_RAG_TIMING] Store '{store.name}' search took {store_search_time:.3f}s - raw={len(results)}, filtered={len(filtered_results)}")
            
            if filtered_results:
                logger.info(f"[GLOBAL_RAG] Store '{store.name}' matched with {len(filtered_results)} results (min_score={min_score})")
                for chunk_id, score in filtered_results[:max_results]:
                    logger.debug(f"[GLOBAL_RAG_DETAIL] Store '{store.name}' - chunk_id={chunk_id}, score={score:.4f}")
                matched_store_names.append(store.name)
                
                # Limit to max_results
                for chunk_id, score in filtered_results[:max_results]:
                    all_semantic_results.append((chunk_id, score, store_id))
            else:
                logger.debug(f"[GLOBAL_RAG] Store '{store.name}' had no semantic matches above threshold {min_score}")
        
        # Convert semantic results to standard format for merging
        if all_semantic_results:
            all_semantic_results.sort(key=lambda x: x[1], reverse=True)
            chunk_ids = [r[0] for r in all_semantic_results]
            scores_map = {r[0]: r[1] for r in all_semantic_results}
            # NC-0.8.0.1.1: Enable document-level keyword filtering for global KB
            semantic_results = await self._get_chunks_by_ids(
                db, chunk_ids, scores_map,
                query=query,
                history=original_query,  # Include original for broader matching
                filter_by_doc_keywords=True
            )
        else:
            semantic_results = []
        
        # Merge semantic and identifier results
        if identifiers and identifier_results:
            final_results = self._merge_search_results(
                semantic_results,
                identifier_results,
                semantic_threshold=global_min_score,
                identifier_boost=0.15,
                top_k=10
            )
            
            # Update matched store names with any identifier-only matches
            for result in final_results:
                ks_id = result.get("knowledge_store_id")
                if ks_id:
                    for store in global_stores:
                        if str(store.id) == ks_id and store.name not in matched_store_names:
                            matched_store_names.append(store.name)
        else:
            final_results = semantic_results
        
        if not final_results:
            total_time = time.time() - total_start
            logger.debug(f"[GLOBAL_RAG_TIMING] No relevant results found in any global store (total time: {total_time:.3f}s)")
            return [], []
        
        total_time = time.time() - total_start
        logger.info(f"[GLOBAL_RAG_TIMING] Total search time: {total_time:.3f}s (embed={embed_time:.3f}s)")
        logger.info(f"[GLOBAL_RAG] Returning {len(final_results)} results from stores: {matched_store_names}")
        
        return final_results, matched_store_names
    
    async def _get_chunks_by_ids(
        self,
        db: AsyncSession,
        chunk_ids: List[str],
        scores: Dict[str, float],
        query: str = "",
        history: str = "",
        filter_by_doc_keywords: bool = False,
    ) -> List[Dict[str, Any]]:
        """
        Fetch chunk details by IDs.
        
        NC-0.8.0.1.1: Optional document-level keyword filtering
        When filter_by_doc_keywords=True, chunks from documents with
        require_keywords_enabled will be filtered based on keyword match.
        """
        result = await db.execute(
            select(DocumentChunk, Document)
            .join(Document, DocumentChunk.document_id == Document.id)
            .where(DocumentChunk.id.in_(chunk_ids))
        )
        rows = result.all()
        
        results = []
        filtered_count = 0
        
        for chunk, document in rows:
            # NC-0.8.0.1.1: Check document-level keyword requirements
            if filter_by_doc_keywords and document.require_keywords_enabled:
                phrases, keywords = _parse_keywords_string(document.required_keywords or "")
                match_mode = document.keyword_match_mode or 'any'
                
                if not _check_keywords_match(query, history, phrases, keywords, match_mode):
                    filtered_count += 1
                    logger.debug(f"[DOC_KEYWORD_FILTER] Document '{document.name}' filtered - keywords not matched (mode={match_mode})")
                    continue
            
            results.append({
                "chunk_id": str(chunk.id),
                "document_id": str(document.id),
                "document_name": document.name,
                "knowledge_store_id": str(document.knowledge_store_id) if document.knowledge_store_id else None,
                "content": chunk.content,
                "chunk_index": chunk.chunk_index,
                "similarity": scores.get(str(chunk.id), 0.0),
                "metadata": chunk.chunk_metadata,
            })
        
        if filtered_count > 0:
            logger.info(f"[DOC_KEYWORD_FILTER] Filtered {filtered_count} chunks due to document keyword requirements")
        
        results.sort(key=lambda x: x["similarity"], reverse=True)
        return results
    
    async def _get_store_document_ids(
        self,
        db: AsyncSession,
        store_ids: List[str],
    ) -> List[str]:
        """Get all document IDs for knowledge stores"""
        result = await db.execute(
            select(Document.id)
            .where(Document.knowledge_store_id.in_(store_ids))
            .where(Document.is_processed == True)
        )
        return [str(row[0]) for row in result.all()]
    
    async def _get_accessible_store_ids(
        self,
        db: AsyncSession,
        user: User,
        store_ids: List[str],
    ) -> List[str]:
        """Get list of knowledge store IDs the user has access to"""
        accessible = []
        
        for store_id in store_ids:
            result = await db.execute(
                select(KnowledgeStore).where(KnowledgeStore.id == store_id)
            )
            store = result.scalar_one_or_none()
            
            if not store:
                continue
            
            if store.owner_id == user.id or store.is_public:
                accessible.append(store_id)
                continue
            
            # Check for explicit share
            result = await db.execute(
                select(KnowledgeStoreShare).where(
                    KnowledgeStoreShare.knowledge_store_id == store_id,
                    KnowledgeStoreShare.shared_with_user_id == user.id
                )
            )
            if result.scalar_one_or_none():
                accessible.append(store_id)
        
        return accessible
    
    async def _keyword_search(
        self,
        db: AsyncSession,
        user: User,
        query: str,
        document_ids: Optional[List[str]],
        top_k: int,
    ) -> List[Dict[str, Any]]:
        """
        Fallback keyword-based search.
        
        When document_ids is not specified, only searches unitemized documents
        (documents not in any knowledge store).
        """
        keywords = query.lower().split()
        
        query_builder = (
            select(DocumentChunk, Document)
            .join(Document, DocumentChunk.document_id == Document.id)
            .where(Document.owner_id == user.id)
            .where(Document.is_processed == True)
        )
        
        if document_ids:
            query_builder = query_builder.where(Document.id.in_(document_ids))
        else:
            # Only search unitemized documents (not in any knowledge store)
            query_builder = query_builder.where(Document.knowledge_store_id.is_(None))
        
        result = await db.execute(query_builder)
        rows = result.all()
        
        results = []
        for chunk, document in rows:
            content_lower = chunk.content.lower()
            score = sum(1 for kw in keywords if kw in content_lower)
            
            if score > 0:
                results.append({
                    "chunk_id": str(chunk.id),
                    "document_id": str(document.id),
                    "document_name": document.name,
                    "content": chunk.content,
                    "chunk_index": chunk.chunk_index,
                    "similarity": score / len(keywords),
                    "metadata": chunk.chunk_metadata,
                })
        
        results.sort(key=lambda x: x["similarity"], reverse=True)
        return results[:top_k]
    
    async def _keyword_search_documents(
        self,
        db: AsyncSession,
        document_ids: List[str],
        query: str,
        top_k: int,
    ) -> List[Dict[str, Any]]:
        """Keyword search across specific documents"""
        keywords = query.lower().split()
        
        result = await db.execute(
            select(DocumentChunk, Document)
            .join(Document, DocumentChunk.document_id == Document.id)
            .where(Document.id.in_(document_ids))
        )
        rows = result.all()
        
        results = []
        for chunk, document in rows:
            content_lower = chunk.content.lower()
            score = sum(1 for kw in keywords if kw in content_lower)
            
            if score > 0:
                results.append({
                    "chunk_id": str(chunk.id),
                    "document_id": str(document.id),
                    "document_name": document.name,
                    "knowledge_store_id": str(document.knowledge_store_id) if document.knowledge_store_id else None,
                    "content": chunk.content,
                    "chunk_index": chunk.chunk_index,
                    "similarity": score / len(keywords),
                    "metadata": chunk.chunk_metadata,
                })
        
        results.sort(key=lambda x: x["similarity"], reverse=True)
        return results[:top_k]
    
    async def get_context_for_query(
        self,
        db: AsyncSession,
        user: User,
        query: str,
        document_ids: Optional[List[str]] = None,
        top_k: int = 5,
        chat_id: Optional[str] = None,
    ) -> str:
        """
        Get formatted context from user's UNITEMIZED documents for LLM prompt.
        This is the main entry point for RAG context retrieval.
        
        Only searches documents that are NOT in any knowledge store.
        Documents in knowledge stores are only accessed when:
        - A custom GPT with that knowledge store is selected (uses get_knowledge_store_context)
        - The user explicitly provides document_ids
        
        This prevents searching subscribed GPTs' knowledge stores when no GPT is selected.
        
        Args:
            chat_id: Optional chat ID for context-aware query enhancement.
        """
        # Enhance query with conversation context (for follow-up questions)
        if chat_id:
            query = await self._enhance_query_with_context(db, query, chat_id)
        
        # If specific document IDs provided, search only those
        if document_ids:
            results = await self.search(db, user, query, document_ids, top_k)
        else:
            # Search user's unitemized documents (not in any knowledge store)
            results = await self.search(db, user, query, None, top_k)
        
        if not results:
            return ""
        
        context_parts = []
        for i, result in enumerate(results, 1):
            context_parts.append(
                f"[Source {i}: {result['document_name']}]\n{result['content']}"
            )
        
        return "\n\n---\n\n".join(context_parts)
    
    async def get_knowledge_store_context(
        self,
        db: AsyncSession,
        user: User,
        query: str,
        knowledge_store_ids: List[str],
        top_k: int = 5,
        bypass_access_check: bool = False,
        chat_id: Optional[str] = None,
        max_context_tokens: Optional[int] = None,
        enable_summarization: bool = True,
    ) -> str:
        """
        Get formatted context from knowledge stores for LLM prompt.
        
        NC-0.8.0.6: Smart RAG Compression
        - Re-ranks results by subject relevance
        - Summarizes chunks if summary is smaller than original
        - Respects token budget
        
        Args:
            bypass_access_check: If True, skip ownership/permission checks.
                                Used when accessing KB through a public GPT.
            chat_id: Optional chat ID for context-aware query enhancement.
            max_context_tokens: Maximum tokens for RAG context (None = no limit)
            enable_summarization: Whether to summarize long chunks
        """
        # Enhance query with conversation context (for follow-up questions)
        if chat_id:
            query = await self._enhance_query_with_context(db, query, chat_id)
        
        results = await self.search_knowledge_stores(
            db, user, query, knowledge_store_ids, top_k,
            bypass_access_check=bypass_access_check,
        )
        
        if not results:
            return ""
        
        debug_enabled = await _is_debug_rag_enabled(db)
        
        # NC-0.8.0.6: Re-rank by subject relevance if we have multiple results
        if len(results) > 1:
            results = await self._rerank_by_subject(query, results, debug_enabled)
        
        # NC-0.8.0.6: Smart compression - summarize if beneficial
        context_parts = []
        total_tokens = 0
        
        for i, result in enumerate(results, 1):
            content = result['content']
            doc_name = result['document_name']
            similarity = result.get('similarity', 0)
            
            # Estimate tokens (rough: 4 chars per token)
            content_tokens = len(content) // 4
            
            # Check token budget
            if max_context_tokens and total_tokens + content_tokens > max_context_tokens:
                if debug_enabled:
                    _log_rag_debug(
                        "Token budget reached, stopping context build",
                        included=i-1,
                        total=len(results),
                        tokens_used=total_tokens
                    )
                break
            
            # NC-0.8.0.6: Summarize if chunk is large and summarization is enabled
            final_content = content
            if enable_summarization and content_tokens > 500:  # Only summarize chunks > ~2000 chars
                summary = await self._summarize_chunk(content, query, debug_enabled)
                if summary:
                    summary_tokens = len(summary) // 4
                    if summary_tokens < content_tokens * 0.7:  # Only use if 30%+ smaller
                        final_content = f"[Summarized]\n{summary}"
                        content_tokens = summary_tokens
                        if debug_enabled:
                            _log_rag_debug(
                                f"Chunk {i} summarized",
                                original_tokens=len(content)//4,
                                summary_tokens=summary_tokens,
                                savings=f"{(1 - summary_tokens/(len(content)//4))*100:.0f}%"
                            )
            
            context_parts.append(
                f"[Source {i}: {doc_name} | Relevance: {similarity:.0%}]\n{final_content}"
            )
            total_tokens += content_tokens
        
        context = "\n\n---\n\n".join(context_parts)
        
        # Log final context length if debug enabled
        if debug_enabled:
            _log_rag_debug(
                "Context built for LLM prompt",
                sources=len(context_parts),
                context_length=len(context),
                estimated_tokens=total_tokens,
                context_preview=context[:200] + "..." if len(context) > 200 else context
            )
        
        return context
    
    async def _rerank_by_subject(
        self,
        query: str,
        results: List[Dict[str, Any]],
        debug_enabled: bool = False,
    ) -> List[Dict[str, Any]]:
        """
        NC-0.8.0.6: Re-rank results by subject relevance.
        
        Uses keyword overlap scoring to boost results that match
        the query subject more closely.
        """
        import re
        
        # Extract key terms from query (simple approach - nouns and important words)
        query_lower = query.lower()
        # Remove common words
        stopwords = {'the', 'a', 'an', 'is', 'are', 'was', 'were', 'what', 'how', 'why', 
                    'when', 'where', 'who', 'which', 'this', 'that', 'these', 'those',
                    'i', 'you', 'he', 'she', 'it', 'we', 'they', 'my', 'your', 'his',
                    'her', 'its', 'our', 'their', 'do', 'does', 'did', 'have', 'has',
                    'had', 'can', 'could', 'will', 'would', 'should', 'may', 'might',
                    'to', 'of', 'in', 'for', 'on', 'with', 'at', 'by', 'from', 'about',
                    'into', 'through', 'during', 'before', 'after', 'above', 'below',
                    'and', 'or', 'but', 'if', 'then', 'than', 'so', 'as', 'be', 'been'}
        
        query_words = set(re.findall(r'\b\w+\b', query_lower)) - stopwords
        
        if not query_words:
            return results  # No meaningful words to match
        
        # Score each result by keyword overlap
        for result in results:
            content_lower = result['content'].lower()
            content_words = set(re.findall(r'\b\w+\b', content_lower))
            
            # Calculate overlap
            overlap = len(query_words & content_words)
            overlap_ratio = overlap / len(query_words) if query_words else 0
            
            # Boost score based on overlap (max 20% boost)
            original_score = result.get('similarity', 0.5)
            boost = overlap_ratio * 0.2
            result['_rerank_score'] = min(1.0, original_score + boost)
            result['_keyword_overlap'] = overlap
        
        # Sort by rerank score
        results.sort(key=lambda x: x.get('_rerank_score', 0), reverse=True)
        
        if debug_enabled:
            _log_rag_debug(
                "Re-ranked results by subject",
                query_keywords=list(query_words)[:10],
                top_overlaps=[r.get('_keyword_overlap', 0) for r in results[:3]]
            )
        
        return results
    
    async def _summarize_chunk(
        self,
        content: str,
        query: str,
        debug_enabled: bool = False,
    ) -> Optional[str]:
        """
        NC-0.8.0.6: Summarize a RAG chunk, focused on the query.
        
        Returns summary if successful, None if summarization fails.
        """
        from app.services.settings_service import SettingsService
        
        try:
            # Get summarization model (use same as main LLM by default)
            # This could be a lighter model for efficiency
            import httpx
            from app.core.config import settings
            
            # Check if summarization is enabled in settings
            # summarize_enabled = await SettingsService.get_bool(db, "rag_summarization_enabled")
            # if not summarize_enabled:
            #     return None
            
            # Simple extractive summarization (no LLM call needed)
            # Extract sentences most relevant to the query
            sentences = re.split(r'(?<=[.!?])\s+', content)
            
            if len(sentences) <= 3:
                return None  # Too short to summarize
            
            query_lower = query.lower()
            query_words = set(re.findall(r'\b\w+\b', query_lower))
            
            # Score sentences by relevance to query
            scored_sentences = []
            for sent in sentences:
                sent_lower = sent.lower()
                sent_words = set(re.findall(r'\b\w+\b', sent_lower))
                overlap = len(query_words & sent_words)
                scored_sentences.append((overlap, sent))
            
            # Sort by relevance and take top sentences
            scored_sentences.sort(key=lambda x: x[0], reverse=True)
            
            # Take top 30% of sentences, minimum 2, maximum 5
            num_sentences = max(2, min(5, len(sentences) // 3))
            top_sentences = [s[1] for s in scored_sentences[:num_sentences]]
            
            # Reorder by original position for coherence
            original_order = {sent: i for i, sent in enumerate(sentences)}
            top_sentences.sort(key=lambda s: original_order.get(s, 999))
            
            summary = ' '.join(top_sentences)
            
            if debug_enabled:
                _log_rag_debug(
                    "Extractive summarization",
                    original_sentences=len(sentences),
                    summary_sentences=len(top_sentences),
                    compression_ratio=f"{len(summary)/len(content)*100:.0f}%"
                )
            
            return summary
            
        except Exception as e:
            if debug_enabled:
                _log_rag_debug(f"Summarization failed: {e}")
            return None
    
    async def rebuild_all_indexes(self, db: AsyncSession) -> Dict[str, int]:
        """Rebuild all knowledge store FAISS indexes"""
        result = await db.execute(select(KnowledgeStore))
        stores = result.scalars().all()
        
        rebuilt = {}
        for store in stores:
            await self._update_knowledge_store_index(db, str(store.id))
            
            # Count vectors in index
            index_manager = self.get_index_manager()
            if str(store.id) in index_manager._indexes:
                rebuilt[str(store.id)] = index_manager._indexes[str(store.id)].ntotal
            else:
                rebuilt[str(store.id)] = 0
        
        return rebuilt


class DocumentProcessor:
    """Process various document types for RAG"""
    
    @staticmethod
    async def extract_text_from_pdf_tika(file_path: str) -> str:
        """Extract text from PDF using Apache Tika"""
        import httpx
        
        try:
            async with httpx.AsyncClient(timeout=60.0) as client:
                with open(file_path, 'rb') as f:
                    pdf_content = f.read()
                
                response = await client.put(
                    settings.TIKA_URL,
                    content=pdf_content,
                    headers={
                        'Content-Type': 'application/pdf',
                        'Accept': 'text/plain',
                    }
                )
                
                if response.status_code == 200:
                    text = response.text.strip()
                    if text:
                        logger.info(f"Tika extracted {len(text)} chars from PDF")
                        return text
                else:
                    logger.warning(f"Tika returned status {response.status_code}")
        except Exception as e:
            logger.warning(f"Tika extraction failed: {e}")
        
        return ""
    
    @staticmethod
    async def extract_text(file_path: str, mime_type: str) -> str:
        """Extract text content from various file types"""
        import json
        path = Path(file_path)
        
        if mime_type in ["text/plain", "text/markdown"]:
            return path.read_text(encoding="utf-8")
        
        elif mime_type == "application/json":
            data = json.loads(path.read_text(encoding="utf-8"))
            return json.dumps(data, indent=2)
        
        elif mime_type == "text/csv":
            import csv
            text_parts = []
            with open(path, "r", encoding="utf-8") as f:
                reader = csv.reader(f)
                for row in reader:
                    text_parts.append(" | ".join(row))
            return "\n".join(text_parts)
        
        elif mime_type == "application/pdf":
            # Try Tika first (better extraction)
            tika_text = await DocumentProcessor.extract_text_from_pdf_tika(file_path)
            if tika_text:
                return tika_text
            
            # Fallback to PyMuPDF
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                text_parts = []
                for page in doc:
                    text_parts.append(page.get_text())
                doc.close()
                return "\n\n".join(text_parts)
            except ImportError:
                logger.warning("PyMuPDF not available for PDF fallback")
                return ""
        
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or path.suffix.lower() == ".docx":
            # Word document (.docx)
            return await DocumentProcessor.extract_text_from_docx(file_path)
        
        elif mime_type == "application/msword" or path.suffix.lower() == ".doc":
            # Legacy Word document (.doc) - try with docx library (may not work for all)
            # For full .doc support, would need antiword or libreoffice
            logger.warning(".doc files may not extract properly - consider converting to .docx")
            return await DocumentProcessor.extract_text_from_docx(file_path)
        
        elif mime_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/vnd.ms-excel"] or path.suffix.lower() in [".xlsx", ".xls"]:
            # Excel spreadsheet
            return await DocumentProcessor.extract_text_from_xlsx(file_path)
        
        elif mime_type in ["application/rtf", "text/rtf"] or path.suffix.lower() == ".rtf":
            # RTF document
            return await DocumentProcessor.extract_text_from_rtf(file_path)
        
        else:
            try:
                return path.read_text(encoding="utf-8")
            except:
                return ""
    
    @staticmethod
    async def extract_text_from_docx(file_path: str) -> str:
        """Extract text from Word document (.docx)"""
        try:
            from docx import Document
            doc = Document(file_path)
            
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(text)} chars from DOCX")
            return text
            
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return ""
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return ""
    
    @staticmethod
    async def extract_text_from_xlsx(file_path: str) -> str:
        """Extract text from Excel spreadsheet (.xlsx/.xls)"""
        try:
            from openpyxl import load_workbook
            
            wb = load_workbook(file_path, data_only=True)
            text_parts = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_parts.append(f"=== Sheet: {sheet_name} ===")
                
                for row in sheet.iter_rows():
                    row_values = []
                    for cell in row:
                        if cell.value is not None:
                            row_values.append(str(cell.value))
                    if row_values:
                        text_parts.append(" | ".join(row_values))
            
            wb.close()
            text = "\n".join(text_parts)
            logger.info(f"Extracted {len(text)} chars from XLSX")
            return text
            
        except ImportError:
            logger.error("openpyxl not installed. Run: pip install openpyxl")
            return ""
        except Exception as e:
            logger.error(f"XLSX extraction failed: {e}")
            # Try with xlrd for .xls files
            if file_path.lower().endswith('.xls'):
                return await DocumentProcessor._extract_text_from_xls_legacy(file_path)
            return ""
    
    @staticmethod
    async def _extract_text_from_xls_legacy(file_path: str) -> str:
        """Extract text from legacy Excel (.xls) using xlrd"""
        try:
            import xlrd
            
            wb = xlrd.open_workbook(file_path)
            text_parts = []
            
            for sheet_idx in range(wb.nsheets):
                sheet = wb.sheet_by_index(sheet_idx)
                text_parts.append(f"=== Sheet: {sheet.name} ===")
                
                for row_idx in range(sheet.nrows):
                    row_values = [str(cell.value) for cell in sheet.row(row_idx) if cell.value]
                    if row_values:
                        text_parts.append(" | ".join(row_values))
            
            text = "\n".join(text_parts)
            logger.info(f"Extracted {len(text)} chars from XLS (legacy)")
            return text
            
        except ImportError:
            logger.error("xlrd not installed. Run: pip install xlrd")
            return ""
        except Exception as e:
            logger.error(f"XLS extraction failed: {e}")
            return ""
    
    @staticmethod
    async def extract_text_from_rtf(file_path: str) -> str:
        """Extract text from RTF document"""
        try:
            from striprtf.striprtf import rtf_to_text
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_content = f.read()
            
            text = rtf_to_text(rtf_content)
            logger.info(f"Extracted {len(text)} chars from RTF")
            return text
            
        except ImportError:
            logger.error("striprtf not installed. Run: pip install striprtf")
            return ""
        except Exception as e:
            logger.error(f"RTF extraction failed: {e}")
            return ""
